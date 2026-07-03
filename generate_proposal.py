"""
generate_proposal.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Generator Proposal CARAKA Desktop v2 — Implementasi Kriptografi
Output : PROPOSAL_CARAKA_v2.docx

Format diadopsi dari "PROPOSAL CARAKA (2).docx":
  • Times New Roman 12pt, justify
  • Margin: L=4cm R=3cm T=3cm B=3cm
  • H1 (BAB)  : 14pt, Bold, Caps, Centered, numbered
  • H2 (SUB BAB): 12pt, Bold, Caps, numbered A/B/C
  • H3        : 12pt, Bold, numbered 1/2/3
  • Line spacing: 1.5x (body), 1.0x (code/table)
  • Tabel sederhana tanpa background warna (akademik)

Dependency: pip install python-docx
Run       : python generate_proposal.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import sys, copy
from pathlib import Path

def ensure_deps():
    import importlib, subprocess
    for pkg, imp in [("python-docx", "docx")]:
        try:
            importlib.import_module(imp)
        except ImportError:
            print(f"[*] Installing {pkg}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])
            print("[OK]")

ensure_deps()

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

TEMPLATE = Path(__file__).parent / "PROPOSAL CARAKA (2).docx"
OUTPUT   = Path(__file__).parent / "PROPOSAL_CARAKA_v2.docx"

# ── Warna minimal untuk tabel header (tetap akademik) ─────────────
HDR_BG  = '2F4F7F'   # dark blue untuk header tabel (satu-satunya warna)
ALT_BG  = 'F2F2F2'   # abu-abu sangat terang untuk baris alt

# ═════════════════════════════════════════════════════════════════
# HELPERS
# ═════════════════════════════════════════════════════════════════

def new_doc():
    """
    Buka template lama untuk mewarisi styles (termasuk numbering BAB),
    lalu hapus semua konten sehingga kita bisa isi ulang.
    """
    if TEMPLATE.exists():
        doc = Document(str(TEMPLATE))
        # Hapus semua body content kecuali elemen terakhir (sectPr)
        body = doc.element.body
        for child in list(body)[:-1]:   # keep last sectPr
            body.remove(child)
        # Reset semua header/footer paragraf agar tidak ada sisa
        for section in doc.sections:
            for hdr in (section.header, section.footer):
                for p in hdr.paragraphs:
                    for run in p.runs:
                        run.text = ''
    else:
        doc = Document()
        print(f"[WARN] Template tidak ditemukan, pakai dokumen kosong")

    # Perbaiki margin sesuai template lama
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(4)
        section.right_margin  = Cm(3)
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(3)

    # Pastikan Normal style: Times New Roman 12pt
    ns = doc.styles['Normal']
    ns.font.name = 'Times New Roman'
    ns.font.size = Pt(12)

    return doc


def cell_shd(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_repeat_header(row):
    """Jadikan baris sebagai header yang berulang di setiap halaman."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)


def body_para(doc, text, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              bold=False, italic=False, size=12, indent_left=0):
    """Paragraf isi utama: TNR 12pt, 1.5 spasi, first-line indent."""
    p   = doc.add_paragraph()
    p.alignment = align
    pf  = p.paragraph_format
    pf.space_before       = Pt(0)
    pf.space_after        = Pt(0)
    pf.line_spacing_rule  = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(1.25)
    if indent_left:
        pf.left_indent = Cm(indent_left)
    run = p.add_run(text)
    run.bold          = bold
    run.italic        = italic
    run.font.name     = 'Times New Roman'
    run.font.size     = Pt(size)
    return p


def caption_para(doc, text):
    """Caption tabel/gambar: TNR 11pt, center, single space."""
    try:
        p = doc.add_paragraph(style='Caption')
    except Exception:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(11)
    run.bold        = True
    return p


def blank(doc, n=1):
    for _ in range(n):
        p  = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing = Pt(12)


def code_block(doc, text, caption=''):
    """Blok kode: Courier New 9pt, single-space, indent kiri."""
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after  = Pt(2)
        r = p.add_run(caption)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.italic    = True

    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent        = Cm(1)
    pf.right_indent       = Cm(1)
    pf.space_before       = Pt(4)
    pf.space_after        = Pt(8)
    pf.line_spacing_rule  = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EFEFEF')
    pPr.append(shd)


def bullet_list(doc, items):
    """Daftar bullet menggunakan style List Paragraph dari template."""
    for item in items:
        try:
            p = doc.add_paragraph(style='List Paragraph')
        except Exception:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.25)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if isinstance(item, tuple):
            label, rest = item
            r1 = p.add_run(f'\u2022  {label}')
            r1.bold = True
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(12)
            r2 = p.add_run(f' {rest}')
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
        else:
            r = p.add_run(f'\u2022  {item}')
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)


def numbered_list(doc, items, start=1):
    """Daftar bernomor sederhana."""
    for i, item in enumerate(items, start):
        p  = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent       = Cm(1.25)
        pf.first_line_indent = Cm(-0.63)
        pf.space_before      = Pt(0)
        pf.space_after       = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if isinstance(item, tuple):
            label, rest = item
            r1 = p.add_run(f'{i}.\t{label}')
            r1.bold      = True
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(12)
            r2 = p.add_run(f' {rest}')
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
        else:
            r = p.add_run(f'{i}.\t{item}')
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)


def simple_table(doc, headers, rows, col_widths=None, caption=''):
    """
    Tabel akademik sederhana:
    - Header: bold, center, background biru gelap, teks putih
    - Baris: TNR 11pt, alternating light grey/white
    - Border: all sides single
    """
    if caption:
        caption_para(doc, caption)

    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hr = tbl.rows[0]
    set_repeat_header(hr)
    for ci, h in enumerate(headers):
        cell = hr.cells[ci]
        cell_shd(cell, HDR_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            cell.width = Cm(col_widths[ci])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after  = Pt(3)
        r = p.add_run(h)
        r.bold          = True
        r.font.name     = 'Times New Roman'
        r.font.size     = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row_data in enumerate(rows):
        tr  = tbl.rows[ri + 1]
        bg  = ALT_BG if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = tr.cells[ci]
            cell_shd(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                cell.width = Cm(col_widths[ci])
            p  = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after  = Pt(2)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

            if isinstance(val, tuple):
                text  = val[0]
                bold  = val[1] if len(val) > 1 else False
                align = val[2] if len(val) > 2 else WD_ALIGN_PARAGRAPH.LEFT
                p.alignment = align
                r = p.add_run(text)
                r.bold      = bold
            else:
                r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

    doc.add_paragraph()   # jarak setelah tabel
    return tbl


def add_h1(doc, text):
    """BAB (Heading 1): centered, bold, caps, 14pt."""
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_h2(doc, text):
    """SUB BAB (Heading 2): left, bold, caps, 12pt."""
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_h3(doc, text):
    """Subsubbab (Heading 3): left, bold, 12pt."""
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


# ═════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════

def cover(doc):
    def ctr(text, size=12, bold=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        r.bold = bold
        return p

    blank(doc, 2)

    ctr('PROPOSAL TUGAS', 16, bold=True)
    ctr('IMPLEMENTASI KRIPTOGRAFI', 16, bold=True, space_after=24)

    blank(doc, 3)

    ctr('CARAKA Desktop: Perancangan dan Implementasi Protokol', 16, bold=True)
    ctr('Komunikasi Mesh Offline Terdesentralisasi Berbasis', 16, bold=True)
    ctr('Lightweight Cryptography (Ascon NIST SP 800-232)', 16, bold=True, space_after=24)

    blank(doc, 3)

    # Penyusun
    ctr('Disusun Oleh:', 12, bold=True, space_after=4)
    blank(doc)
    ctr('Andika Aryansyach Fauzan (2322101878)', 14, bold=True)
    ctr('Mahendra Nur Hidayat (23221019)', 14, bold=True)
    ctr('Rafi Putra Fadlurrahman (23221963)', 14, bold=True)

    blank(doc, 3)

    ctr('Rekayasa Sistem Kriptografi', 14, bold=True)
    ctr('Politeknik Siber dan Sandi Negara', 14, bold=True)
    ctr('2026', 14, bold=True)

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# DAFTAR ISI (manual)
# ═════════════════════════════════════════════════════════════════

def daftar_isi(doc):
    add_h1(doc, 'DAFTAR ISI')
    blank(doc)

    entries = [
        ('DAFTAR ISI',                                        'ii'),
        ('DAFTAR GAMBAR',                                     'iii'),
        ('DAFTAR TABEL',                                      'iv'),
        ('BAB I PENDAHULUAN',                                 '1'),
        ('    A. Latar Belakang',                             '1'),
        ('    B. Identifikasi Masalah',                       '3'),
        ('    C. Rumusan Masalah',                            '3'),
        ('    D. Tujuan Penelitian',                          '4'),
        ('    E. Manfaat Penelitian',                         '4'),
        ('BAB II TINJAUAN PUSTAKA',                           '5'),
        ('    A. Komparasi dengan Sistem Serupa',             '5'),
        ('    B. Analisis Algoritma Lightweight Cryptography','7'),
        ('BAB III PERANCANGAN SISTEM',                        '12'),
        ('    A. Arsitektur Berlapis CARAKA Desktop',         '12'),
        ('    B. Protokol CLAMP v0.1',                        '14'),
        ('    C. Model Kunci dan Manajemen Vault',            '18'),
        ('BAB IV IMPLEMENTASI',                               '20'),
        ('    A. Fitur Inti',                                 '20'),
        ('    B. Fitur Keamanan (v2 Baru)',                   '21'),
        ('    C. Fitur Lanjutan (v2 Baru)',                   '22'),
        ('    D. Status Implementasi',                        '23'),
        ('BAB V ARSITEKTUR KEAMANAN DAN THREAT MODEL',        '24'),
        ('    A. Properti Keamanan yang Dijamin',             '24'),
        ('    B. Model Ancaman',                              '25'),
        ('BAB VI RENCANA EVALUASI',                           '27'),
        ('    A. Microbenchmark Kriptografi',                 '27'),
        ('    B. Network-Level Benchmark',                    '28'),
        ('    C. Hipotesis',                                  '29'),
        ('BAB VII PENUTUP',                                   '30'),
        ('    A. Kesimpulan Rancangan',                       '30'),
        ('    B. Jadwal Pelaksanaan',                         '30'),
        ('    C. Luaran yang Diharapkan',                     '31'),
        ('DAFTAR PUSTAKA',                                    '33'),
    ]

    # Tabel sederhana untuk TOC
    tbl = doc.add_table(rows=len(entries), cols=2)
    tbl.style = 'Table Grid'
    # Hapus border dengan membuat border transparan
    for ri, (title, page) in enumerate(entries):
        c1, c2 = tbl.rows[ri].cells
        c1.width = Cm(12)
        c2.width = Cm(2)
        for cell in (c1, c2):
            # Hilangkan border
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
            tcPr.append(tcBorders)

        p1 = c1.paragraphs[0]
        pf1 = p1.paragraph_format
        pf1.space_before = Pt(0)
        pf1.space_after  = Pt(0)
        pf1.line_spacing_rule = WD_LINE_SPACING.SINGLE
        bold_flag = not title.startswith('    ')
        r1 = p1.add_run(title)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r1.bold = bold_flag

        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r2 = p2.add_run(page)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# DAFTAR GAMBAR & TABEL (placeholder)
# ═════════════════════════════════════════════════════════════════

def daftar_gambar(doc):
    add_h1(doc, 'DAFTAR GAMBAR')
    blank(doc)
    gambar_entries = [
        ('Gambar 1.', 'Arsitektur Berlapis CARAKA Desktop'),
        ('Gambar 2.', 'Urutan Implementasi Per Fase (Dependency Graph)'),
        ('Gambar 3.', 'Format Biner Paket CLAMP v0.1'),
        ('Gambar 4.', 'Model Hierarki Kunci CARAKA Desktop'),
        ('Gambar 5.', 'Sequence Diagram Pengiriman DM E2EE'),
        ('Gambar 6.', 'Diagram Epidemic Store-and-Forward Sync'),
    ]
    for label, keterangan in gambar_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f'{label} ')
        r1.bold = True; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
        r2 = p.add_run(keterangan)
        r2.font.name='Times New Roman'; r2.font.size=Pt(12)
    doc.add_page_break()


def daftar_tabel(doc):
    add_h1(doc, 'DAFTAR TABEL')
    blank(doc)
    tabel_entries = [
        ('Tabel 1.', 'Perbandingan Sistem Komunikasi Offline yang Ada'),
        ('Tabel 2.', 'Matriks Perbandingan Algoritma LWC Finalis NIST'),
        ('Tabel 3.', 'Stack Kriptografi CARAKA Desktop'),
        ('Tabel 4.', 'Spesifikasi Lingkungan Target'),
        ('Tabel 5.', 'Model Ancaman dan Mitigasi CARAKA Desktop'),
        ('Tabel 6.', 'Status Implementasi Per Fase'),
        ('Tabel 7.', 'Parameter Microbenchmark Kriptografi'),
        ('Tabel 8.', 'Parameter Network-Level Benchmark'),
        ('Tabel 9.', 'Jadwal Pelaksanaan Penelitian'),
        ('Tabel 10.','Luaran yang Diharapkan'),
        ('Tabel 11.','Stack Teknologi Final CARAKA Desktop v2'),
    ]
    for label, keterangan in tabel_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f'{label} ')
        r1.bold = True; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
        r2 = p.add_run(keterangan)
        r2.font.name='Times New Roman'; r2.font.size=Pt(12)
    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# BAB I — PENDAHULUAN
# ═════════════════════════════════════════════════════════════════

def bab1(doc):
    add_h1(doc, 'BAB I\nPENDAHULUAN')

    add_h2(doc, 'Latar Belakang')
    body_para(doc,
        'Infrastruktur komunikasi digital kontemporer memiliki dependensi yang hampir mutlak '
        'terhadap entitas terpusat: server, Content Delivery Network (CDN), penyedia layanan '
        'internet, dan tulang punggung jaringan global. Dependensi ini menciptakan single point '
        'of failure yang secara arsitektural berbahaya dalam tiga skenario utama, yaitu: (1) '
        'kegagalan infrastruktur akibat bencana alam; (2) pembatasan akses oleh pihak berwenang; '
        'dan (3) degradasi layanan pada kondisi darurat massal.')
    body_para(doc,
        'Jaringan Peer-to-Peer (P2P) Mesh menawarkan solusi arsitektural alternatif di mana '
        'setiap perangkat berperan ganda sebagai klien sekaligus router, memungkinkan pesan '
        'mencapai tujuannya melalui jalur alternatif tanpa bergantung pada entitas pusat. '
        'Paradigma ini sangat relevan untuk koordinasi darurat bencana, komunikasi di daerah '
        'terpencil tanpa akses internet, serta kebutuhan privasi komunikasi yang tinggi.')
    body_para(doc,
        'Pada tahun 2023, National Institute of Standards and Technology (NIST) mengumumkan '
        'terpilihnya keluarga Ascon sebagai standar Lightweight Cryptography (LWC), yang kemudian '
        'dipublikasikan sebagai NIST Special Publication 800-232 pada Agustus 2025. Standar ini '
        'dirancang khusus untuk perangkat dengan keterbatasan sumber daya, namun tidak terdapat '
        'implementasi yang mengintegrasikan standar ini secara penuh ke dalam protokol komunikasi '
        'mesh offline terdesentralisasi di platform desktop.')
    body_para(doc,
        'Mengamankan komunikasi pada arsitektur multi-hop menghadirkan tiga tantangan kriptografi '
        'yang belum terpecahkan secara optimal. Pertama, masalah overhead ukuran paket: pada radio '
        'low-bandwidth seperti LoRa dengan batas 256 byte per paket, tanda tangan Ed25519 '
        'mengkonsumsi 64 byte (lebih dari 25% kapasitas paket) hanya untuk autentikasi. Kedua, '
        'kebocoran metadata routing: sistem yang ada membiarkan routing header dalam plaintext '
        'sehingga pihak ketiga dapat memetakan topologi jaringan secara pasif. Ketiga, manajemen '
        'kunci pada jaringan delay-tolerant: Perfect Forward Secrecy berbasis Double Ratchet '
        'memerlukan urutan pesan yang terjamin, sebuah asumsi yang tidak berlaku pada jaringan '
        'offline dengan penundaan tak tentu.')
    body_para(doc,
        'Berdasarkan identifikasi masalah tersebut, proyek CARAKA Desktop (Cryptographically '
        'Authenticated Relay Architecture for Knowledge and Autonomy) diusulkan sebagai solusi '
        'komprehensif. CARAKA Desktop mengimplementasikan protokol CLAMP (Compact Lightweight '
        'Authenticated Mesh Protocol) yang menggunakan Ascon sebagai satu-satunya keluarga '
        'kriptografi untuk AEAD, hash, dan MAC berlapis-hop, dengan inovasi utama menggantikan '
        'Ed25519 per-paket (64 byte) dengan Ascon-MAC per-hop (16 byte), menghasilkan penghematan '
        '75% overhead autentikasi.')
    blank(doc)

    add_h2(doc, 'Identifikasi Masalah')
    body_para(doc,
        'Berdasarkan latar belakang yang telah dipaparkan, identifikasi masalah dalam penelitian '
        'ini adalah sebagai berikut:')
    bullet_list(doc, [
        'Tidak ada sistem komunikasi mesh offline yang mengintegrasikan standar NIST LWC (Ascon '
        'NIST SP 800-232) ke dalam desain protokol keamanannya.',
        'Overhead autentikasi per-paket pada sistem yang ada (64 byte Ed25519) terlalu besar '
        'untuk jaringan radio low-bandwidth seperti LoRa (256 byte/paket).',
        'Routing header yang dikirim dalam plaintext mengekspos metadata topologi jaringan '
        'kepada pihak ketiga yang melakukan traffic analysis pasif.',
        'Mekanisme sinkronisasi store-and-forward yang ada tidak bersifat privacy-preserving '
        'karena node perantara dapat mengakses konten pesan.',
    ])
    blank(doc)

    add_h2(doc, 'Rumusan Masalah')
    body_para(doc, 'Berdasarkan identifikasi masalah di atas, rumusan masalah penelitian ini adalah:')
    numbered_list(doc, [
        'Bagaimana merancang protokol komunikasi mesh terdesentralisasi berbasis Ascon (NIST SP '
        '800-232) yang meminimalkan cryptographic overhead per-hop untuk autentikasi relay?',
        'Bagaimana mengimplementasikan skema autentikasi multi-hop menggunakan Ascon-MAC sebagai '
        'pengganti Ed25519 per-paket, dan seberapa besar reduksi overhead yang dicapai secara empiris?',
        'Bagaimana performa Ascon-AEAD128 dibandingkan AES-256-GCM dalam parameter encryption '
        'throughput, overhead paket, end-to-end latency, dan memory footprint pada platform desktop?',
        'Bagaimana merancang mekanisme sinkronisasi store-and-forward yang privacy-preserving '
        'menggunakan hash fingerprint tanpa mengekspos konten kepada node perantara?',
    ])
    blank(doc)

    add_h2(doc, 'Tujuan Penelitian')
    body_para(doc, 'Tujuan dari penelitian ini adalah:')
    numbered_list(doc, [
        'Merancang dan mengimplementasikan protokol CLAMP yang mengintegrasikan Ascon sebagai '
        'satu-satunya keluarga kriptografi untuk AEAD, hash, dan MAC dalam mesh networking.',
        'Mengukur dan membuktikan reduksi overhead autentikasi dari penggantian Ed25519 (64 byte) '
        'dengan Ascon-MAC (16 byte) secara empiris.',
        'Mengevaluasi performa Ascon-AEAD128 dibandingkan algoritma konvensional pada konteks '
        'desktop mesh networking.',
        'Merancang mekanisme Epidemic Sync berbasis SHA-256 fingerprint vector yang tidak '
        'mengekspos konten pesan kepada node perantara.',
    ])
    blank(doc)

    add_h2(doc, 'Manfaat Penelitian')
    body_para(doc,
        'Penelitian ini diharapkan memberikan manfaat sebagai berikut. Dari sisi akademik, '
        'penelitian ini menghasilkan bukti empiris pertama evaluasi Ascon NIST SP 800-232 dalam '
        'konteks desktop mesh networking, melengkapi literatur yang selama ini hanya mengevaluasi '
        'LWC pada perangkat tertanam. Dari sisi praktis, penelitian ini menghasilkan prototipe '
        'aplikasi komunikasi offline-first yang dapat digunakan untuk koordinasi darurat bencana '
        'dan komunikasi di daerah tanpa infrastruktur internet.')

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# BAB II — TINJAUAN PUSTAKA
# ═════════════════════════════════════════════════════════════════

def bab2(doc):
    add_h1(doc, 'BAB II\nTINJAUAN PUSTAKA')

    add_h2(doc, 'Komparasi dengan Sistem Serupa')
    body_para(doc,
        'Sistem komunikasi offline yang ada menghadapi dilema mendasar yang belum terpecahkan '
        'antara keamanan dan efisiensi jaringan. Berikut adalah tinjauan kritis terhadap empat '
        'sistem paling relevan yang telah ada di pasaran.')

    add_h3(doc, 'Meshtastic')
    body_para(doc,
        'Meshtastic merupakan platform mesh off-grid berbasis radio LoRa pada perangkat ESP32 '
        'dan NRF52. Sejak firmware v2.5, Meshtastic menggunakan AES-256-CTR dengan Pre-Shared '
        'Key untuk channel dan AES-CCM dengan X25519 untuk Direct Messages. Meskipun demikian, '
        'terdapat beberapa keterbatasan kritis: (1) routing header dikirim dalam plaintext '
        'sehingga pihak ketiga dapat memetakan topologi jaringan secara pasif; (2) tanda tangan '
        'XEdDSA berbasis Ed25519 mengkonsumsi 64 byte per paket, lebih dari 25% kapasitas LoRa '
        '(256 byte); dan (3) Pre-Shared Key statis berarti kompromi satu node membocorkan '
        'seluruh riwayat channel.')

    add_h3(doc, 'Berty / Wesh Protocol')
    body_para(doc,
        'Berty adalah protokol komunikasi offline-first terdesentralisasi di atas IPFS dan '
        'OrbitDB yang menggunakan X25519 untuk enkripsi dan Ed25519 untuk tanda tangan per-event. '
        'Berty secara eksplisit mengakui bahwa protokolnya belum siap untuk data sensitif tinggi '
        'dan belum sepenuhnya diaudit. Overhead Ed25519 per-event yang direplikasi melalui gossip '
        'juga menghasilkan beban signifikan pada jaringan low-bandwidth.')

    add_h3(doc, 'Briar')
    body_para(doc,
        'Briar adalah aplikasi pesan P2P yang dirancang untuk aktivis dan jurnalis, '
        'mengimplementasikan konstruksi mirip Signal dengan Perfect Forward Secrecy berbasis '
        'Double Ratchet dan AES-GCM. Keterbatasan utamanya adalah routing offline yang bergantung '
        'pada social graph sehingga pesan hanya dapat diteruskan melalui kontak bersama, bukan '
        'jaringan mesh generik.')

    add_h3(doc, 'Secure Scuttlebutt (SSB)')
    body_para(doc,
        'Secure Scuttlebutt adalah protokol gossip berbasis append-only signed log per identitas '
        'di mana setiap entri log ditandatangani dengan Ed25519. Penggunaan log permanen '
        'menciptakan rekam jejak digital yang bertentangan fundamental dengan kebutuhan privasi '
        'pesan sementara, dan sinkronisasi yang memerlukan seluruh riwayat log tidak efisien '
        'untuk komunikasi pesan biasa.')
    blank(doc)

    # Tabel perbandingan
    simple_table(doc,
        headers=['Sistem', 'Overhead Auth', 'Keamanan Metadata', 'Mode Offline', 'Primitif Kriptografi'],
        rows=[
            ['Meshtastic', '64B Ed25519\n(per paket)', 'Tidak (header\nplaintext)', 'Penuh', 'AES-CTR/AES-CCM'],
            ['Berty', 'Ed25519\nper-event', 'Tidak (gossip\nterbuka)', 'Terbatas', 'X25519 + Ed25519'],
            ['Briar', 'AES-GCM +\nDouble Ratchet', 'Ya (via Tor,\nbutuh internet)', 'Terbatas\n(social graph)', 'Signal Protocol'],
            ['Secure Scuttlebutt', 'Ed25519\nper-log', 'Tidak (log\npermanen)', 'Penuh', 'Ed25519'],
            [('CARAKA Desktop\n(Proposal)', True, WD_ALIGN_PARAGRAPH.CENTER),
             ('17B Ascon-MAC\n(per hop)', True, WD_ALIGN_PARAGRAPH.CENTER),
             ('Ya (MAC per-hop\nterenkripsi)', True, WD_ALIGN_PARAGRAPH.CENTER),
             ('Penuh', True, WD_ALIGN_PARAGRAPH.CENTER),
             ('Ascon NIST SP\n800-232', True, WD_ALIGN_PARAGRAPH.CENTER)],
        ],
        col_widths=[3.5, 3, 3.5, 2.5, 3.5],
        caption='Tabel 1. Perbandingan Sistem Komunikasi Offline yang Ada'
    )

    body_para(doc,
        'Berdasarkan komparasi di atas, terdapat research gap yang signifikan: tidak ada sistem '
        'komunikasi mesh offline yang mengintegrasikan standar NIST LWC (Ascon) ke dalam desain '
        'protokol keamanannya untuk mengatasi masalah cryptographic overhead pada autentikasi '
        'multi-hop relay dan enkripsi metadata routing.')
    blank(doc)

    add_h2(doc, 'Analisis Algoritma Lightweight Cryptography')
    body_para(doc,
        'NIST memulai proses standardisasi Lightweight Cryptography pada tahun 2013 untuk '
        'mengatasi ketidaksesuaian antara algoritma kriptografi standar (AES-GCM, SHA-256) dengan '
        'keterbatasan perangkat tertanam. Setelah satu dekade evaluasi ketat — termasuk kompetisi '
        'CAESAR dan proses NIST LWC resmi — keluarga Ascon terpilih pada Februari 2023 dan '
        'dipublikasikan sebagai NIST Special Publication 800-232 pada Agustus 2025.')

    # Matriks perbandingan
    simple_table(doc,
        headers=['Kriteria', 'ASCON\n(Dipilih)', 'TinyJambu', 'GIFT-COFB', 'Grain-128', 'PRESENT'],
        rows=[
            ['Status NIST', 'Standar\nSP 800-232', 'Finalis,\ntidak dipilih', 'Finalis,\ntidak dipilih', 'Finalis,\ntidak dipilih', 'ISO/IEC\n29192-2'],
            ['Tingkat\nKeamanan', '128-bit\npenuh', '~120-bit\n(bermasalah)', '~64-bit\nefektif', '128-bit,\nkeystream terbatas', '80/128-bit\n(80 usang)'],
            ['Kapabilitas', 'AEAD +\nHash + XOF', 'AEAD\nsaja', 'AEAD\nsaja', 'AEAD\nsaja', 'Block Cipher\nsaja'],
            ['Isu\nKriptanalisis', 'Minimal', 'Birthday-bound\nslide attacks', '64-bit tag\nforgery', 'Batas panjang\nkeystream', 'Sweet32\n(64-bit blok)'],
            ['Dukungan\nRust', 'Matang\n(RustCrypto)', 'Tidak ada\ncrate resmi', 'Tidak ada\ncrate resmi', 'Terbatas', 'Minimal'],
            ['Cocok\nDesktop Mesh', 'Ya -\nSangat Cocok', 'Tidak', 'Terbatas', 'Dengan syarat\nketat', 'Tidak'],
        ],
        col_widths=[3, 3, 2.7, 2.7, 2.7, 2.7],
        caption='Tabel 2. Matriks Perbandingan Algoritma LWC Finalis NIST'
    )

    add_h3(doc, 'Ascon (Algoritma Terpilih)')
    body_para(doc,
        'Ascon adalah sponge-based permutation 320-bit yang mencakup: Ascon-AEAD128 (kunci 128-bit, '
        'nonce 128-bit, tag 128-bit), Ascon-Hash256 (fungsi hash 256-bit), dan Ascon-XOF128 '
        '(Extendable Output Function untuk KDF). State internal 320-bit yang dioperasikan dalam '
        'mode duplex memungkinkan satu keluarga primitif memenuhi seluruh kebutuhan AEAD, hashing, '
        'dan KDF. Ascon dipilih karena: (1) satu-satunya kandidat memenuhi seluruh kriteria keamanan '
        '128-bit; (2) merupakan standar NIST aktif; (3) memiliki kapabilitas AEAD+Hash+XOF dalam '
        'satu keluarga; dan (4) ekosistem Rust paling matang melalui crate ascon-aead di RustCrypto.')

    add_h3(doc, 'TinyJambu (Ditolak)')
    body_para(doc,
        'TinyJambu mengalami serangkaian temuan kriptanalisis berupa birthday-bound slide attacks '
        'pada permutasi kunci P2 yang membuktikan asumsi ideal permutation dalam bukti keamanan '
        'mode AEAD-nya tidak berlaku, dengan security margin kurang dari 8 bit terhadap kompleksitas '
        'data. Tidak tersedia crate Rust yang resmi dan terawat.')

    add_h3(doc, 'GIFT-COFB (Ditolak)')
    body_para(doc,
        'Analisis oleh Khairallah (IACR ePrint 2021/648) menunjukkan bahwa batas keamanan efektif '
        'GIFT-COFB berperilaku seperti AEAD dengan tag 64-bit (bukan 128-bit) dalam skenario '
        'high-forgery count pada jaringan long-lived. Selain itu tidak tersedia crate Rust.')

    add_h3(doc, 'Grain-128AEAD (Kondisional)')
    body_para(doc,
        'Grain-128AEAD membatasi total panjang keystream per pasangan kunci/IV hingga sekitar '
        '2^80 bit. Pada jaringan mesh aktif dengan banyak pesan, manajemen rotasi kunci yang '
        'ketat menjadi beban operasional yang tidak realistis.')

    add_h3(doc, 'PRESENT (Ditolak)')
    body_para(doc,
        'PRESENT adalah block cipher 64-bit yang rentan terhadap serangan kolisi birthday bound '
        '(Sweet32) untuk volume data besar. Versi kunci 80-bit telah dianggap usang dan tidak '
        'memiliki mode AEAD bawaan.')
    blank(doc)

    # Stack tabel
    simple_table(doc,
        headers=['Fungsi', 'Primitif', 'Crate Rust'],
        rows=[
            ['Enkripsi Payload (E2EE)', 'Ascon-AEAD128 (NIST SP 800-232)', 'ascon-aead v0.4'],
            ['Fungsi Hash / Fingerprint', 'SHA-256', 'sha2 v0.10'],
            ['MAC per-Hop Relay', 'HKDF-HMAC-SHA256 (keyed, 16B tag)', 'hkdf v0.12'],
            ['Key Exchange (ECDH)', 'X25519 Diffie-Hellman', 'x25519-dalek v2'],
            ['Key Derivation (KDF)', 'HKDF-SHA256 (RFC 5869)', 'hkdf + sha2'],
            ['Penyimpanan Kunci', 'Argon2id Vault (m=32MiB, t=2)', 'argon2 v0.5'],
            ['Transport Anonim', 'Tor Hidden Service', 'arti-client'],
            ['Serialisasi Paket', 'Binary (Bincode)', 'bincode v1'],
        ],
        col_widths=[5, 6, 5],
        caption='Tabel 3. Stack Kriptografi CARAKA Desktop'
    )

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# BAB III — PERANCANGAN SISTEM
# ═════════════════════════════════════════════════════════════════

def bab3(doc):
    add_h1(doc, 'BAB III\nPERANCANGAN SISTEM')

    add_h2(doc, 'Arsitektur Berlapis CARAKA Desktop')
    body_para(doc,
        'CARAKA Desktop dirancang dengan arsitektur berlapis (layered architecture) yang memisahkan '
        'concern kriptografi, protokol, jaringan, penyimpanan, dan antarmuka pengguna. Arsitektur '
        'ini memastikan setiap lapisan hanya bergantung pada lapisan di bawahnya, menjamin modularitas '
        'dan kemudahan pengujian per komponen.')

    code_block(doc, """\
+------------------------------------------------------------------+
|         PRESENTATION LAYER -- Tauri v2 (HTML/CSS/JS)            |
|      Chat UI | Radar Topology | Settings | Emergency Mode        |
+------------------------------------------------------------------+
|         APPLICATION LAYER -- commands.rs (Tauri IPC)            |
|   Vault | Message | File Transfer | Tor | Emergency Commands     |
+------------------------+-----------------------------------------+
|   PROTOCOL LAYER       |       STORAGE LAYER                     |
|   packet.rs (CLAMP)    |  store.rs -- SQLite (ciphertext only)   |
|   routing.rs           |  sync.rs  -- Epidemic Sync              |
|   (Flooding+Trust+RL)  |  (Fingerprint Vector Privacy-pres.)     |
+------------------------+-----------------------------------------+
|   NETWORK LAYER        |       CRYPTOGRAPHY LAYER                |
|   discovery.rs UDP7770 |  crypto.rs  -- Ascon-AEAD128            |
|   transport.rs TCP7771 |  keys.rs    -- X25519 ECDH + HKDF       |
|   tor.rs (Onion Svc)   |  vault.rs   -- Argon2id Vault           |
|   hotspot.rs (netsh)   |  network_monitor.rs                     |
+------------------------+-----------------------------------------+""",
        'Gambar 1. Arsitektur Berlapis CARAKA Desktop')

    body_para(doc,
        'Komponen kriptografi (Ascon-AEAD128, X25519, Argon2id) menempati lapisan terbawah dan '
        'menjadi fondasi yang digunakan oleh semua lapisan di atasnya. Lapisan jaringan '
        'menangani komunikasi UDP/TCP dan transport opsional via Tor. Lapisan protokol '
        'mengimplementasikan CLAMP dengan mekanisme routing, autentikasi hop, dan deduplication '
        'paket. Lapisan penyimpanan menggunakan SQLite yang hanya menyimpan ciphertext, '
        'sedangkan lapisan aplikasi menyediakan antarmuka IPC antara frontend Tauri dan backend Rust.')

    add_h3(doc, 'Urutan Implementasi')
    body_para(doc,
        'Implementasi dilakukan secara berurutan sesuai dependency graph untuk memastikan '
        'setiap modul dapat diuji secara independen sebelum modul yang bergantung padanya '
        'diimplementasikan:')

    code_block(doc, """\
Fase 1: crypto.rs + keys.rs + vault.rs  --> Fondasi kriptografi
  |
Fase 2: packet.rs                        --> Format biner CLAMP
  |
Fase 3: routing.rs                       --> Flooding + Trust Score + Rate Limiter
  |
Fase 4: discovery.rs + transport.rs      --> UDP discovery + TCP transport
  |
Fase 5: store.rs + sync.rs               --> SQLite + Epidemic Sync
  |
Fase 6: commands.rs + GUI + Tor + Hotspot + File Transfer
  |
Fase 7: Benchmarking (Criterion.rs + Network)""",
        'Gambar 2. Urutan Implementasi Per Fase (Dependency Graph)')

    simple_table(doc,
        headers=['Parameter', 'Spesifikasi'],
        rows=[
            ['Platform Target', 'Windows 10/11 (64-bit) dan Ubuntu 22.04 LTS'],
            ['Transport Utama', 'UDP :7770 (peer discovery) dan TCP :7771 (data + sync)'],
            ['Transport Anonim', 'Tor onion service via arti-client (opsional)'],
            ['Skala Jaringan', '2 sampai 15 node dalam satu segmen LAN/Wi-Fi'],
            ['Topologi', 'Ad-hoc mesh tanpa hierarki tetap (decentralized)'],
            ['Bahasa Backend', 'Rust 1.78+ stable (memory safety tanpa garbage collector)'],
            ['Framework GUI', 'Tauri v2 (cross-platform desktop application)'],
            ['Konektivitas', 'Offline-first -- tidak memerlukan koneksi internet'],
        ],
        col_widths=[5, 11],
        caption='Tabel 4. Spesifikasi Lingkungan Target'
    )
    blank(doc)

    add_h2(doc, 'Protokol CLAMP v0.1')
    body_para(doc,
        'CLAMP (Compact Lightweight Authenticated Mesh Protocol) adalah protokol lapisan aplikasi '
        'biner yang dirancang khusus untuk CARAKA Desktop. Setiap paket memiliki fixed overhead '
        '62 byte yang mencakup routing header (13 byte), hop authentication (17 byte), nonce '
        '(16 byte), dan AEAD tag (16 byte).')

    add_h3(doc, 'Struktur Paket Biner CLAMP')
    code_block(doc, """\
Offset  Size  Field           Deskripsi
------  ----  --------------  --------------------------------------------------
  0       2   magic           0xCA, 0x52  -- identifikasi protokol CARAKA
  2       1   version         0x01        -- versi protokol
  3       1   packet_type     0x01=DM | 0x02=Channel | 0x03=SyncReq
                               0x04=SyncResp | 0x05=Hello | 0x06=SyncData
                               0x07=Broadcast (Emergency)
  4       1   ttl             0-7, dikurangi setiap relay hop
  5       8   packet_id       origin_node_id[0..4] || OsRng[4..8]
             =================== 13 byte ROUTING HEADER ===================
 13       1   hop_counter     0 saat origin, +1 per relay
 14      16   hop_mac_tag     HKDF-HMAC(ch_key, pkt_id||hop_ctr||relay[0..4])
             ================ 17 byte HOP AUTHENTICATION ==================
 30      16   nonce           timestamp_u32_LE[0..4] || OsRng[4..16]
 46       N   ciphertext      Ascon-AEAD128 encrypted inner payload
 46+N    16   aead_tag        Authentication tag 128-bit
             ============= Total fixed overhead: 62 byte ==================""",
        'Gambar 3. Format Biner Paket CLAMP v0.1')

    body_para(doc,
        'Inovasi utama CLAMP terletak pada lapisan Hop Authentication: setiap relay node '
        'memvalidasi Hop-MAC menggunakan kunci channel yang dibagi out-of-band, memperbarui '
        'hop_counter, dan menghitung ulang MAC sebelum meneruskan paket. Paket dengan MAC tidak '
        'valid di-drop secara diam-diam tanpa respons (silent drop), mencegah adversary mendapatkan '
        'informasi tentang keberadaan node di jaringan.')

    add_h3(doc, 'Perbandingan Overhead CLAMP vs. Meshtastic')
    code_block(doc, """\
Meshtastic (paket LoRa 256 byte):
  Header + Routing  : ~30 byte  (PLAINTEXT -- mengekspos topologi)
  AES-CCM nonce+tag :  16 byte
  XEdDSA / Ed25519  :  64 byte  (tanda tangan per-paket)
  Total overhead    : ~110 byte  --> hanya ~146 byte untuk payload

CARAKA CLAMP (paket LoRa 256 byte):
  Routing Header    : 13 byte
  Hop Auth (MAC)    : 17 byte   (hop_counter + 16 byte HMAC tag)
  Ascon Nonce       : 16 byte
  Ascon AEAD Tag    : 16 byte
  Total overhead    :  62 byte  --> ~194 byte untuk payload (+33% lebih besar)

  Penghematan overhead autentikasi:
    64 byte (Ed25519) --> 16 byte (Ascon-MAC) = HEMAT 75%""",
        'Perbandingan overhead CARAKA CLAMP vs. Meshtastic per paket LoRa 256 byte')

    add_h3(doc, 'Alur Pengiriman Direct Message E2EE')
    code_block(doc, """\
Alice (Pengirim)         Bob (Relay)            Charlie (Penerima)
      |                        |                        |
  [1] Derive DM-Key            |                        |
      shared = X25519(alice_priv, charlie_pub)           |
      dm_key = HKDF(shared, session_id, msg_counter)     |
      |                        |                        |
  [2] Enkripsi Payload         |                        |
      nonce = ts[0..4] || rand                           |
      (ct, tag) = Ascon-AEAD128(dm_key, nonce,           |
                                plaintext, aad=header)   |
      |                        |                        |
  [3] Build CLAMP Packet       |                        |
      TTL=7, hop_counter=0     |                        |
      hop_mac = HMAC(ch_key, pkt_id||0||alice[0..4])    |
      |                        |                        |
      +------ CLAMP Packet --->|                        |
                          [4] Validasi Hop-MAC           |
                              OK  --> TTL--, hop_ctr++   |
                                   Hitung ulang MAC      |
                              FAIL --> DROP diam-diam    |
                              +------- Packet ---------->|
                                                    [5] Validasi Hop-MAC OK
                                                        Derive dm_key (mirror)
                                                        Ascon-AEAD128.decrypt
                                                        --> Tampilkan pesan""",
        'Gambar 5. Sequence Diagram Pengiriman DM E2EE Melalui Relay')

    add_h3(doc, 'Epidemic Store-and-Forward Sync')
    code_block(doc, """\
Ketika Node A dan Node B pertama kali terhubung:

  A ---[ SyncReq: Fingerprint Vector A ]--> B
        Vector = [SHA256(pkt_id_1), SHA256(pkt_id_2), ...]

  B ---[ SyncResp: Fingerprint Vector B ]--> A

  Kedua node menghitung selisih vektor
  --> identifikasi pesan yang belum ada

  A ---[ SyncData: ciphertext missing di B ]--> B

PROPERTI PRIVASI:
  * Node perantara hanya melihat ciphertext (tidak dapat membaca isi)
  * SHA-256 fingerprint tidak mengekspos konten pesan
  * Ascon-AEAD128 tag memastikan ciphertext tidak dapat dimodifikasi""",
        'Gambar 6. Diagram Epidemic Store-and-Forward Sync')
    blank(doc)

    add_h2(doc, 'Model Kunci dan Manajemen Vault')
    body_para(doc,
        'Berbeda dari proposal v1 yang menggunakan Windows Credential Manager, CARAKA v2 '
        'mengimplementasikan vault file terenkripsi yang portabel lintas platform dengan format '
        '80 byte. Private key disimpan terenkripsi menggunakan Argon2id sebagai KDF dan '
        'Ascon-AEAD128 sebagai algoritma enkripsi.')

    code_block(doc, """\
Hierarki Kunci CARAKA Desktop:

  passphrase --> Argon2id(m=32768KiB, t=2, p=1) --> 16B Ascon key
                                                      |
                     vault.key (80 byte file)          |
                     [0..16]  Argon2id salt (OsRng)    |
                     [16..32] Ascon-AEAD nonce (OsRng) |
                     [32..80] ciphertext + AEAD tag     v
                                                 Ascon-AEAD128.encrypt(
                                                   key=kdf_output,
                                                   plaintext=x25519_private_key
                                                 )

  X25519 Private Key (diproteksi vault di atas)
    |
    | ECDH: shared = X25519(my_private_key, peer_public_key)
    v
  HKDF-SHA256(ikm=shared, info="CARAKA-DM-v1"||my_id||peer_id||
                                session_id||msg_counter, len=16)
    |
    +--> DM-Key (16B) --> Ascon-AEAD128 enkripsi payload
    +--> Session ID unik per sesi --> simple forward secrecy""",
        'Gambar 4. Model Hierarki Kunci CARAKA Desktop')

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# BAB IV — IMPLEMENTASI
# ═════════════════════════════════════════════════════════════════

def bab4(doc):
    add_h1(doc, 'BAB IV\nIMPLEMENTASI')

    add_h2(doc, 'Fitur Inti')
    body_para(doc,
        'Fitur inti merupakan komponen dasar CARAKA Desktop yang diimplementasikan pada Fase 1 '
        'hingga Fase 4 dari rencana implementasi. Fitur-fitur ini membentuk fondasi fungsional '
        'sistem komunikasi mesh offline.')
    bullet_list(doc, [
        ('End-to-End Encryption (E2EE): ',
         'Setiap Direct Message dienkripsi Ascon-AEAD128 dengan kunci unik hasil ECDH X25519 '
         'sebelum meninggalkan perangkat. Node relay tidak pernah dapat membaca isi pesan.'),
        ('Protokol CLAMP (Mesh Networking): ',
         'Controlled Flooding dengan TTL maksimum 7 hop, Packet ID LRU-512 untuk deduplication, '
         'Trust Score per-peer, dan Token Bucket rate limiter (burst 200 paket, 100 paket/detik).'),
        ('Auto Peer Discovery: ',
         'UDP Broadcast per 30 detik ke seluruh antarmuka jaringan aktif. Mendukung penambahan '
         'peer manual via IP:Port, QR Code, atau Invite Code base64url (format caraka0/caraka1).'),
        ('Offline-First + Epidemic Sync: ',
         'Pesan tersimpan terenkripsi di SQLite. Epidemic Sync via SHA-256 Fingerprint Vector '
         'dilakukan otomatis saat peer kembali terhubung tanpa mengekspos konten ke relay.'),
    ])
    blank(doc)

    add_h2(doc, 'Fitur Keamanan (v2 Baru)')
    body_para(doc,
        'Fitur keamanan baru diimplementasikan pada Fase 5 sebagai penguatan terhadap model '
        'keamanan CARAKA Desktop, mencakup perbaikan signifikan dari proposal v1.')
    bullet_list(doc, [
        ('Argon2id Vault [BARU]: ',
         'Private key disimpan dalam vault file terenkripsi (80 byte, portabel lintas platform) '
         'menggunakan Argon2id (m=32768 KiB, t=2) sebagai KDF dan Ascon-AEAD128 sebagai cipher. '
         'Menggantikan Windows Credential Manager yang tidak portabel dari proposal v1.'),
        ('QR Code + Safety Number [BARU]: ',
         'Verifikasi identitas peer via QR code yang berisi Node ID (X25519 Public Key). Safety '
         'Number berbasis SHA-256 canonical untuk verifikasi fingerprint secara out-of-band.'),
        ('Token Bucket Rate Limiter [BARU]: ',
         'Perlindungan per-peer: burst 200 paket, 100 paket/detik. Pelanggaran secara otomatis '
         'menurunkan Trust Score untuk mitigasi Sybil Attack dan network flooding.'),
        ('Text Invite Code [BARU]: ',
         'Kode undangan base64url (caraka0 untuk LAN, caraka1 untuk Tor) berisi Node ID dan '
         'alamat koneksi, memungkinkan sharing peer via teks biasa tanpa scanner QR code.'),
    ])
    blank(doc)

    add_h2(doc, 'Fitur Lanjutan (v2 Baru)')
    body_para(doc,
        'Fitur lanjutan diimplementasikan pada Fase 6 dan memperluas kemampuan CARAKA Desktop '
        'untuk skenario darurat dan kebutuhan privasi tingkat lanjut.')
    bullet_list(doc, [
        ('Emergency Broadcast [BARU]: ',
         'Siaran darurat bertipe INFO/EVAC/STATUS/RESOURCE dengan flooding mesh TTL=7. '
         'Payload tidak dienkripsi untuk memastikan semua node dapat membaca pesan darurat.'),
        ('Emergency Hotspot Mode [BARU]: ',
         'Mengaktifkan Windows Mobile Hotspot dengan SSID "CARAKA-Emergency" (open network) '
         'via netsh wlan. Subnet scan 192.168.137.x untuk mendeteksi peer tanpa infrastruktur router.'),
        ('Tor Hidden Service [BARU]: ',
         'Setiap node mendapatkan alamat .onion persisten via arti-client (pure-Rust Tor). '
         'Invite Code format caraka1 untuk koneksi via Tor dengan anonimitas penuh.'),
        ('File Transfer E2EE [BARU]: ',
         'Transfer file hingga 5 MB per file melalui jalur enkripsi DM yang sama (Ascon-AEAD128 '
         'dan ECDH X25519). File di-chunk dan dienkripsi sebelum dikirim melalui jaringan mesh.'),
        ('Reply-to Message [BARU]: ',
         'Direct Message mendukung field reply_to_id dan reply_to_text untuk thread pesan. '
         'Quote pesan yang dibalas disertakan dalam payload terenkripsi.'),
        ('CI/CD Release Pipeline [BARU]: ',
         'GitHub Actions secara otomatis membangun binary Tauri untuk Windows (NSIS installer '
         'dan MSI) serta mempublikasikannya ke GitHub Releases saat tag v*.*.* di-push.'),
    ])
    blank(doc)

    add_h2(doc, 'Status Implementasi')
    simple_table(doc,
        headers=['Fase', 'Nama Fase', 'Deliverable Utama', 'Status'],
        rows=[
            ['Fase 1', 'Core Crypto Engine', 'crypto.rs, keys.rs, vault.rs, unit test', 'Selesai'],
            ['Fase 2', 'CLAMP Engine', 'packet.rs: binary framing, LRU-512, semua packet types', 'Selesai'],
            ['Fase 3', 'P2P Networking', 'discovery.rs, transport.rs, routing.rs (trust+rate limiter)', 'Selesai'],
            ['Fase 4', 'GUI Tauri v2', 'index.html, main.js: chat, radar, peer list, settings', 'Selesai'],
            ['Fase 5', 'Security Hardening', 'Argon2id vault, QR verification, Safety Number, rate limiting', 'Selesai'],
            ['Fase 6', 'Fitur Lanjutan', 'Emergency Mode, Tor, File Transfer, Reply-to, CI/CD', 'Selesai'],
            ['Fase 7', 'Evaluasi', 'Criterion.rs benchmark, network test multi-topologi, makalah', 'Dalam Progress'],
        ],
        col_widths=[1.5, 3.5, 7.5, 3.5],
        caption='Tabel 6. Status Implementasi Per Fase'
    )

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# BAB V — KEAMANAN DAN THREAT MODEL
# ═════════════════════════════════════════════════════════════════

def bab5(doc):
    add_h1(doc, 'BAB V\nARSITEKTUR KEAMANAN DAN THREAT MODEL')

    add_h2(doc, 'Properti Keamanan yang Dijamin')
    body_para(doc,
        'Sistem CARAKA Desktop dirancang untuk menjamin enam properti keamanan utama yang '
        'secara kolektif memberikan perlindungan menyeluruh terhadap ancaman yang diidentifikasi '
        'dalam model ancaman.')
    numbered_list(doc, [
        ('Confidentiality (E2EE): ',
         'Setiap pesan dienkripsi Ascon-AEAD128 dengan kunci unik per-sesi sebelum meninggalkan '
         'perangkat pengirim. Node relay tidak dapat membaca isi pesan bahkan saat proses Epidemic Sync.'),
        ('Integrity & Authentication: ',
         'Autentikasi beroperasi pada dua lapisan: Hop-MAC per relay node dan AEAD tag 128-bit '
         'end-to-end. Modifikasi ciphertext apapun menyebabkan dekripsi gagal secara deterministik.'),
        ('Replay Protection: ',
         'LRU cache 512 Packet ID dikombinasikan dengan timestamp window ±300 detik mencegah '
         'pengiriman ulang paket lama secara efektif dengan kompleksitas lookup O(1).'),
        ('Best-Effort Forward Secrecy: ',
         'Session ID dan msg_counter yang unik per pesan memberikan lapisan forward secrecy '
         'sederhana tanpa memerlukan ratchet yang membutuhkan sinkronisasi urutan pesan ketat.'),
        ('Key Protection: ',
         'Private key dilindungi Argon2id vault dengan parameter m=32768 KiB (resistensi GPU/ASIC) '
         'dan tidak pernah tersimpan sebagai plaintext di disk dalam kondisi apapun.'),
        ('Transport Privacy (Opsional): ',
         'Tor Hidden Service via arti-client memberikan anonimitas transport untuk pengguna '
         'yang membutuhkan perlindungan metadata lebih kuat dari traffic analysis.'),
    ])
    blank(doc)

    add_h2(doc, 'Model Ancaman')
    body_para(doc,
        'CARAKA dirancang menghadapi adversary yang memiliki kemampuan: (1) mencegat seluruh '
        'lalu lintas jaringan secara pasif; (2) menyuntikkan, mengubah, atau memblokir paket '
        'secara aktif; dan (3) menjalankan node berbahaya untuk serangan Sybil. Adversary '
        'diasumsikan TIDAK dapat membobol primitif kriptografi yang digunakan atau mengakses '
        'perangkat secara fisik.')

    simple_table(doc,
        headers=['Kategori', 'Ancaman Spesifik', 'Level Risiko', 'Mitigasi dalam CARAKA'],
        rows=[
            ['Confidentiality', 'Penyadapan payload (Eavesdropping)', 'TINGGI',
             'E2EE Ascon-AEAD128; ciphertext tidak terbaca tanpa kunci DM'],
            ['Integrity', 'Modifikasi ciphertext (Tampering)', 'TINGGI',
             'AEAD Tag 128-bit; modifikasi apapun menyebabkan dekripsi gagal'],
            ['Replay', 'Pengiriman ulang paket lama', 'SEDANG',
             'LRU-512 Packet ID cache + timestamp window ±300 detik'],
            ['Authentication', 'Penyamaran identitas (Impersonation)', 'SEDANG',
             'TOFU + QR Code verification + Safety Number (SHA-256)'],
            ['Availability', 'Network Flooding / DoS', 'SEDANG',
             'TTL max 7 hop + Token Bucket (100 pkt/det) + Trust Score'],
            ['Availability', 'Sybil Attack', 'SEDANG',
             'Trust Score berbasis perilaku; node baru dimulai dengan skor rendah'],
            ['Privacy', 'Traffic Analysis (metadata)', 'SEDANG',
             'Payload E2EE; header minimal; transport anonim via Tor tersedia'],
            ['Key Security', 'Pencurian private key dari disk', 'SEDANG',
             'Argon2id vault: key dienkripsi Ascon-AEAD128, tidak pernah plaintext'],
            ['Routing', 'Manipulasi routing header', 'RENDAH',
             'Hop-MAC divalidasi setiap relay; paket invalid di-drop diam-diam'],
            ['Physical', 'Kompromi perangkat fisik', 'Di luar cakupan',
             'Vault passphrase-protected; akses fisik tetap berbahaya'],
            ['Quantum', 'Serangan Post-Quantum', 'Di luar cakupan v0.1',
             'X25519 tidak PQC-resistant; Kyber KEM direncanakan untuk v0.3'],
        ],
        col_widths=[3, 4, 2.5, 6.5],
        caption='Tabel 5. Model Ancaman dan Mitigasi CARAKA Desktop'
    )

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# BAB VI — RENCANA EVALUASI
# ═════════════════════════════════════════════════════════════════

def bab6(doc):
    add_h1(doc, 'BAB VI\nRENCANA EVALUASI')

    body_para(doc,
        'Evaluasi dilakukan pada dua level yang komplementer untuk membuktikan nilai akademik '
        'protokol yang diusulkan: microbenchmark kriptografi untuk mengukur performa primitif '
        'kriptografi secara terisolasi, dan network-level benchmark untuk mengukur performa '
        'sistem secara end-to-end pada topologi mesh yang dikontrol.')

    add_h2(doc, 'Microbenchmark Kriptografi (Criterion.rs)')
    body_para(doc,
        'Microbenchmark bertujuan mengukur dan membandingkan performa Ascon-AEAD128 terhadap '
        'AES-256-GCM dan ChaCha20-Poly1305 pada platform desktop menggunakan framework Rust '
        'Criterion yang menghasilkan analisis statistik rigorosa.')

    simple_table(doc,
        headers=['Metrik', 'Satuan', 'Kondisi Uji'],
        rows=[
            ['Encryption Throughput', 'MB/s', 'Ukuran pesan: 64B, 256B, 1KB, 4KB, 16KB'],
            ['Decryption Throughput', 'MB/s', 'Sama dengan kondisi enkripsi'],
            ['Waktu Komputasi MAC', 'mikrodetik', 'HKDF-HMAC-SHA256 untuk input 13 byte (CLAMP header)'],
            ['Waktu Derivasi Kunci', 'mikrodetik', 'HKDF-SHA256: cold start vs. cached'],
            ['Waktu Buka Vault', 'milidetik', 'Argon2id KDF: cold start vs. subsequent unlock'],
            ['Memory Footprint', 'KB', 'Peak RAM usage per operasi kriptografi'],
        ],
        col_widths=[4.5, 2.5, 9],
        caption='Tabel 7. Parameter Microbenchmark Kriptografi'
    )

    body_para(doc,
        'Hipotesis yang akan dibuktikan: Ascon-AEAD128 memiliki throughput lebih rendah dari '
        'AES-256-GCM pada CPU dengan akselerasi hardware AES-NI, namun lebih tinggi pada CPU '
        'tanpa akselerasi hardware, sesuai desain LWC yang bersifat hardware-agnostic.')
    blank(doc)

    add_h2(doc, 'Network-Level Benchmark (Multi-Node LAN)')
    body_para(doc,
        'Network-level benchmark bertujuan mengukur performa sistem secara end-to-end pada '
        'topologi jaringan yang dikontrol. Pengujian dilakukan dengan tiga konfigurasi topologi: '
        '(1) Linear A→B→C→D→E (5 hop, worst-case latency), (2) Star dengan 1 hub dan 4 spoke, '
        'dan (3) Mesh acak dengan 10 node dan rata-rata 3 tetangga per node.')

    simple_table(doc,
        headers=['Metrik', 'Satuan', 'Keterangan'],
        rows=[
            ['End-to-End Latency', 'milidetik', 'Dari pengiriman hingga penerimaan, per konfigurasi topologi'],
            ['Message Delivery Ratio', 'persen', 'Dari 1.000 pesan yang dikirim, berapa yang berhasil diterima'],
            ['Ciphertext Expansion', 'persen', '(ukuran paket CARAKA / ukuran plaintext) x 100 persen'],
            ['Hop Overhead per Node', 'byte', 'Overhead tambahan yang ditambahkan relay per hop'],
            ['Sync Throughput', 'pesan/detik', 'Kecepatan Epidemic Sync antara dua node (1.000 pesan)'],
        ],
        col_widths=[4, 2.5, 9.5],
        caption='Tabel 8. Parameter Network-Level Benchmark'
    )
    blank(doc)

    add_h2(doc, 'Hipotesis Penelitian')
    body_para(doc, 'Empat hipotesis berikut akan dibuktikan secara empiris dalam evaluasi:')
    numbered_list(doc, [
        ('H1: ',
         'Overhead autentikasi CARAKA (17 byte) adalah 73% lebih kecil dari Meshtastic (64 byte), '
         'menghasilkan 33% ruang payload lebih besar untuk paket LoRa 256 byte.'),
        ('H2: ',
         'Ascon-AEAD128 memiliki throughput lebih rendah dari AES-256-GCM pada CPU dengan '
         'AES-NI, namun lebih tinggi pada CPU tanpa akselerasi hardware.'),
        ('H3: ',
         'Message Delivery Ratio lebih dari atau sama dengan 95% pada topologi linear 5-hop '
         'untuk 1.000 pesan di jaringan LAN normal.'),
        ('H4: ',
         'Epidemic Sync tidak mengekspos plaintext kepada node perantara, diverifikasi melalui '
         'Wireshark packet capture yang hanya menampilkan ciphertext.'),
    ])

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# BAB VII — PENUTUP
# ═════════════════════════════════════════════════════════════════

def bab7(doc):
    add_h1(doc, 'BAB VII\nPENUTUP')

    add_h2(doc, 'Kesimpulan Rancangan')
    body_para(doc,
        'Proposal ini memaparkan rancangan dan implementasi CARAKA Desktop sebagai sistem '
        'komunikasi mesh offline-first terdesentralisasi yang mengintegrasikan standar NIST '
        'Lightweight Cryptography (Ascon NIST SP 800-232) secara penuh. Protokol CLAMP v0.1 '
        'yang dirancang menghasilkan penghematan 75% overhead autentikasi dibandingkan sistem '
        'sejenis (Ed25519 64 byte vs. Ascon-MAC 16 byte) sambil mempertahankan keamanan MAC '
        '128-bit yang setara.')
    body_para(doc,
        'Keseluruhan 6 fase implementasi telah berhasil diselesaikan, mencakup: fondasi kriptografi '
        '(Ascon-AEAD128, X25519, Argon2id Vault), protokol CLAMP dengan mesh routing, GUI desktop '
        'berbasis Tauri v2, fitur keamanan lanjutan (QR verification, rate limiting, Safety Number), '
        'dan fitur Emergency Mode (Hotspot, Broadcast, Tor, File Transfer). Saat ini sedang berjalan '
        'Fase 7 untuk evaluasi empiris melalui microbenchmark dan network-level benchmark.')
    blank(doc)

    add_h2(doc, 'Jadwal Pelaksanaan')
    simple_table(doc,
        headers=['Minggu', 'Kegiatan', 'Deliverable', 'Status'],
        rows=[
            ['1', 'Studi literatur, setup environment, desain awal', 'Ringkasan literatur, environment berjalan', 'Selesai'],
            ['2', 'Implementasi modul kriptografi (Ascon, MAC, Hash)', 'crypto.rs + unit test lengkap', 'Selesai'],
            ['3', 'Argon2id vault + manajemen kunci X25519', 'vault.rs + keys.rs update', 'Selesai'],
            ['4', 'CLAMP binary framing semua packet types', 'packet.rs + unit test', 'Selesai'],
            ['5', 'Peer discovery + Token Bucket rate limiter', 'discovery.rs + routing.rs', 'Selesai'],
            ['6', 'TCP transport + Epidemic Sync', 'transport.rs + sync.rs + store.rs', 'Selesai'],
            ['7', 'GUI Tauri v2, QR Code, Safety Number', 'Aplikasi desktop berjalan', 'Selesai'],
            ['8', 'Emergency Mode: Hotspot, Broadcast, Tor, File Transfer', 'hotspot.rs + tor.rs + file_transfer.rs', 'Selesai'],
            ['9', 'Microbenchmark kriptografi (Criterion.rs)', 'Dataset benchmark (CSV/HTML Criterion)', 'Dalam Progress'],
            ['10', 'Network benchmark multi-topologi + makalah final', 'Laporan evaluasi + makalah akademik', 'Belum Dimulai'],
        ],
        col_widths=[1.5, 5, 6, 3.5],
        caption='Tabel 9. Jadwal Pelaksanaan Penelitian CARAKA Desktop'
    )
    blank(doc)

    add_h2(doc, 'Luaran yang Diharapkan')
    simple_table(doc,
        headers=['No.', 'Luaran', 'Deskripsi', 'Status'],
        rows=[
            ['1', 'Prototipe CARAKA Desktop',
             'Aplikasi desktop fungsional (Windows/Linux): GUI Tauri v2, E2EE DM, '
             'Emergency Mode, Tor, File Transfer. Binary di GitHub Releases.',
             'Tercapai'],
            ['2', 'Spesifikasi CLAMP v0.1',
             'Dokumen teknis formal (docs/CLAMP-SPEC.md) mendefinisikan format paket, '
             'operasi kriptografi, semantik routing, dan model kunci.',
             'Tercapai'],
            ['3', 'Repository Open-Source',
             'Kode sumber di github.com/0xAre/CARAKA-DEKSTOP dengan dokumentasi '
             'teknis lengkap dan CI/CD pipeline otomatis.',
             'Tercapai'],
            ['4', 'Dataset Benchmark',
             'Data evaluasi empiris Ascon vs. AES-GCM vs. ChaCha20 dalam format '
             'CSV + Criterion HTML report (open data).',
             'Dalam Progress'],
            ['5', 'Makalah Akademik',
             'Laporan dengan format publikasi: tinjauan pustaka, desain CLAMP, '
             'hasil evaluasi kuantitatif, analisis hipotesis, kesimpulan.',
             'Belum Dimulai'],
        ],
        col_widths=[0.7, 3.5, 9, 2.8],
        caption='Tabel 10. Luaran yang Diharapkan'
    )
    blank(doc)

    simple_table(doc,
        headers=['Lapisan', 'Teknologi', 'Versi', 'Fungsi'],
        rows=[
            ['Framework Desktop', 'Tauri', 'v2.x', 'Cross-platform; backend Rust native'],
            ['Backend', 'Rust', '1.78+', 'Memory safety tanpa garbage collector'],
            ['Frontend', 'HTML5 + CSS3 + Vanilla JS', '--', 'Minimalisasi frontend dependencies'],
            ['Enkripsi Simetris', 'Ascon-AEAD128 (ascon-aead)', '0.4.x', 'NIST SP 800-232 LWC Standard'],
            ['Key Exchange', 'X25519 (x25519-dalek)', '2.x', 'ECDH Diffie-Hellman; diaudit formal'],
            ['Key Derivation', 'HKDF-SHA256 (hkdf + sha2)', 'latest', 'Standar RFC 5869'],
            ['Vault', 'Argon2id (argon2)', '0.5.x', 'm=32MiB, t=2; portabel lintas platform'],
            ['Database', 'SQLite (rusqlite)', '0.39+', 'Hanya menyimpan ciphertext'],
            ['Transport Anonim', 'Tor (arti-client)', 'latest', 'Pure-Rust Tor; onion service persisten'],
            ['Async Runtime', 'Tokio', '1.x', 'De-facto async runtime Rust'],
            ['Benchmarking', 'Criterion.rs', '0.5.x', 'Statistical benchmarking framework'],
        ],
        col_widths=[3.5, 4, 2, 6.5],
        caption='Tabel 11. Stack Teknologi Final CARAKA Desktop v2'
    )

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════
# DAFTAR PUSTAKA
# ═════════════════════════════════════════════════════════════════

def daftar_pustaka(doc):
    add_h1(doc, 'DAFTAR PUSTAKA')
    blank(doc)

    refs = [
        'Meshtastic Project. (2025). Updated Security Implementation. '
        'https://meshtastic.org/docs/development/reference/encryption-technical/',

        'Berty Technologies. (2024). Wesh Protocol Technical Documentation. '
        'https://berty.tech/docs/protocol/',

        'Berty Technologies. (2024). Challenges in Building a Distributed Messaging System. '
        'https://berty.tech/challenges',

        'Briar Project. (2024). How Briar Works. https://briarproject.org/how-it-works/',

        'Tarr, D., Lavoie, C., Meyer, A., & Kermarrec, A.-M. (2019). Secure Scuttlebutt: An '
        'Identity-Centric Protocol for Subjective and Decentralized Applications. Proceedings '
        'of ACM ICN 2019. https://doi.org/10.1145/3357150.3357396',

        'Meshtastic Project. (2025). Known Limitations of Meshtastic Encryption. '
        'https://meshtastic.org/docs/about/overview/encryption/limitations/',

        'NIST. (2023). Lightweight Cryptography Project. '
        'https://csrc.nist.gov/Projects/lightweight-cryptography',

        'NIST. (2025). NIST Special Publication 800-232: Ascon-Based Lightweight Cryptography '
        'Standards for Constrained Devices.',

        'Dobraunig, C., Eichlseder, M., Mendel, F., & Schlaffer, M. (2021). Ascon v1.2: '
        'Lightweight Authenticated Encryption and Hashing. IACR ePrint 2021/1574.',

        'Saha, D., et al. (2022). Birthday-Bound Slide Attacks on TinyJAMBUs '
        'Keyed-Permutations. Proceedings of ASIACRYPT 2022.',

        'Banik, S., et al. (2017). GIFT: A Small PRESENT -- Towards Reaching the Limit of '
        'Lightweight Encryption. IACR ePrint 2017/622.',

        'Khairallah, M. (2021). Security of COFB against Chosen Ciphertext Attacks. '
        'IACR ePrint 2021/648.',

        'Bogdanov, A., et al. (2007). PRESENT: An Ultra-Lightweight Block Cipher. '
        'Proceedings of CHES 2007. https://doi.org/10.1007/978-3-540-74735-2_31',

        'Bhargavan, K., & Leurent, G. (2016). On the Practical (In-)Security of 64-bit Block '
        'Ciphers: Collision Attacks on HTTP over TLS and OpenVPN (Sweet32). ACM CCS 2016.',

        'Banik, S., et al. (2023). NIST IR 8454: Status Report on the Final Round of the '
        'NIST Lightweight Cryptography Standardization Process. NIST.',

        'Biryukov, A., & Khovratovich, D. (2016). Argon2: New Generation of Memory-Hard '
        'Functions for Password Hashing and Other Applications. IEEE EuroS&P 2016.',

        'RustCrypto Project. (2024). ascon-aead: Pure Rust implementation of Ascon. '
        'https://github.com/RustCrypto/AEADs/tree/master/ascon-aead',

        'The Tor Project / Arti Contributors. (2024). arti-client: A pure-Rust Tor client. '
        'https://gitlab.torproject.org/tpo/core/arti',
    ]

    for i, ref in enumerate(refs, 1):
        p  = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent         = Cm(1.5)
        pf.first_line_indent   = Cm(-1.5)
        pf.space_before        = Pt(0)
        pf.space_after         = Pt(0)
        pf.line_spacing_rule   = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f'[{i}]  ')
        r1.bold           = True
        r1.font.name      = 'Times New Roman'
        r1.font.size      = Pt(12)
        r2 = p.add_run(ref)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)


# ═════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════

def add_footer(doc):
    """Tambahkan nomor halaman di tengah footer."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Field PAGE
        run = fp.add_run()
        fldBegin   = OxmlElement('w:fldChar')
        fldBegin.set(qn('w:fldCharType'), 'begin')
        instrText  = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldSep     = OxmlElement('w:fldChar')
        fldSep.set(qn('w:fldCharType'), 'separate')
        fldEnd     = OxmlElement('w:fldChar')
        fldEnd.set(qn('w:fldCharType'), 'end')

        run._r.append(fldBegin)
        run._r.append(instrText)
        run._r.append(fldSep)
        run._r.append(fldEnd)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def main():
    print('=' * 62)
    print('  CARAKA Desktop -- Proposal Generator v2.0')
    print('  Format: PROPOSAL CARAKA (2).docx (Times New Roman, A4)')
    print('=' * 62)

    doc = new_doc()
    add_footer(doc)

    steps = [
        ('Halaman Sampul',   cover),
        ('Daftar Isi',       daftar_isi),
        ('Daftar Gambar',    daftar_gambar),
        ('Daftar Tabel',     daftar_tabel),
        ('BAB I',            bab1),
        ('BAB II',           bab2),
        ('BAB III',          bab3),
        ('BAB IV',           bab4),
        ('BAB V',            bab5),
        ('BAB VI',           bab6),
        ('BAB VII',          bab7),
        ('Daftar Pustaka',   daftar_pustaka),
    ]

    for i, (label, fn) in enumerate(steps, 1):
        print(f'  [{i:02d}/{len(steps)}] Membangun {label}...')
        fn(doc)

    doc.save(str(OUTPUT))
    size = OUTPUT.stat().st_size
    print()
    print('=' * 62)
    print(f'  OK  Berhasil disimpan!')
    print(f'  Dokumen : {OUTPUT.name}')
    print(f'  Lokasi  : {OUTPUT.parent}')
    print(f'  Ukuran  : {size // 1024} KB ({size:,} bytes)')
    print('=' * 62)
    print()
    print('  Buka dengan Microsoft Word atau LibreOffice Writer.')
    print('  Isi [Nama Dosen] dan bagian bracket lainnya sesuai kebutuhan.')


if __name__ == '__main__':
    main()
