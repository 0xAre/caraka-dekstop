from docx import Document
from docx.oxml.ns import qn
from lxml import etree

doc = Document(r'e:\Project APP\CARAKA-DEKSTOP\PROPOSAL CARAKA (2).docx')

# Heading 2 and 3
for lvl in [2, 3]:
    h = doc.styles[f'Heading {lvl}']
    print(f'=== H{lvl} XML ===')
    print(etree.tostring(h._element, pretty_print=True).decode()[:2000])

# Normal style
n = doc.styles['Normal']
print('=== Normal XML ===')
print(etree.tostring(n._element, pretty_print=True).decode()[:2000])

# Table 0 inspect
print()
print('=== TABLE 0 ===')
t = doc.tables[0]
print(f'rows={len(t.rows)}, cols={len(t.columns)}')
for ri, row in enumerate(t.rows[:3]):
    for ci, cell in enumerate(row.cells[:4]):
        txt = cell.text[:40]
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        bg = '?'
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            key = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'
            if shd is not None:
                bg = shd.get(key, '?')
        p = cell.paragraphs[0]
        run_info = ''
        if p.runs:
            r = p.runs[0]
            try:
                clr = str(r.font.color.rgb)
            except Exception:
                clr = 'None'
            run_info = f'bold={r.bold} sz={r.font.size} name={r.font.name} color={clr}'
        print(f'  [{ri}][{ci}] bg={bg!r} aln={p.alignment} txt={txt!r}')
        print(f'         {run_info}')

# Caption style
print()
print('=== Caption style ===')
try:
    cap = doc.styles['Caption']
    print(etree.tostring(cap._element, pretty_print=True).decode()[:1500])
except Exception as e:
    print(e)

# List Paragraph style
print()
print('=== List Paragraph ===')
try:
    lp = doc.styles['List Paragraph']
    print(etree.tostring(lp._element, pretty_print=True).decode()[:1500])
except Exception as e:
    print(e)
